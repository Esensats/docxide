"""Layout Engine and Render Context — Visitor pattern for VDOM rendering."""

from docx import Document as NativeDocument
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .styles import StyleToken, StyleSheet
from .elements import Cell


class RenderContext:
    """Controlled gateway to the native backend."""
    def __init__(self, native_doc: NativeDocument, style_sheet: StyleSheet):
        self._doc = native_doc
        self._style_sheet = style_sheet
        self._current_section = None
        self._page_break_pending = False

    @property
    def native_doc(self) -> NativeDocument:
        return self._doc

    @property
    def style_sheet(self) -> StyleSheet:
        return self._style_sheet

    def add_paragraph(self, style_token: StyleToken) -> 'NativeParagraph':
        """Add a paragraph with proper formatting."""
        para = self._doc.add_paragraph()
        self._apply_paragraph_style(para, style_token)
        return para

    def add_run(self, paragraph, text: str, style_token: StyleToken):
        """Add a run to a paragraph with proper formatting."""
        run = paragraph.add_run(text)
        self._apply_run_style(run, style_token)
        return run

    def _apply_paragraph_style(self, para, token: StyleToken):
        """Apply paragraph-level formatting from a style token."""
        pf = para.paragraph_format

        # Alignment
        if token.alignment is not None:
            pf.alignment = token.alignment

        # Indentation
        if token.indent_first_line is not None:
            pf.first_line_indent = Cm(token.indent_first_line)
        else:
            pf.first_line_indent = Cm(0)

        if token.indent_left > 0:
            pf.left_indent = Cm(token.indent_left)
        if token.indent_right > 0:
            pf.right_indent = Cm(token.indent_right)

        # Spacing
        pf.space_before = Pt(token.space_before)
        pf.space_after = Pt(token.space_after)

        # Line spacing
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.line_spacing = token.line_spacing

        # Page break before
        if token.page_break_before:
            pf.page_break_before = True

    def _apply_run_style(self, run, token: StyleToken):
        """Apply character-level formatting from a style token."""
        font = run.font
        font.name = token.font_name
        font.size = Pt(token.font_size)
        font.bold = token.bold
        font.italic = token.italic

        if token.color:
            font.color.rgb = RGBColor.from_string(token.color)

        # East Asian font
        rpr = run._r.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.insert(0, rfonts)
        rfonts.set(qn('w:ascii'), token.font_name)
        rfonts.set(qn('w:hAnsi'), token.font_name)
        rfonts.set(qn('w:eastAsia'), token.font_name)


class LayoutEngine:
    """Visitor that renders the VDOM tree to a native docx."""
    def __init__(self, style_sheet: StyleSheet):
        self.style_sheet = style_sheet
        self._native_doc = NativeDocument()
        self._section_counter = 0

    def render(self, root, filepath: str = None):
        """Render the VDOM tree to a docx file."""
        ctx = RenderContext(self._native_doc, self.style_sheet)

        # Set page dimensions on first section
        self._setup_page_dimensions(ctx)

        # Render all children (Sections)
        for child in root.children:
            child.render(self, ctx)

        # Validate output
        self._validate()

        if filepath:
            self._native_doc.save(filepath)

        return self._native_doc

    def _setup_page_dimensions(self, ctx: RenderContext):
        """Set page width/height on the default first section."""
        section = ctx.native_doc.sections[0]
        section.page_width = Cm(self.style_sheet.page_width_cm)
        section.page_height = Cm(self.style_sheet.page_height_cm)

    def _validate(self):
        """Check for common issues in the generated document."""
        issues = []
        for i, para in enumerate(self._native_doc.paragraphs):
            for run in para.runs:
                if 'object at 0x' in run.text:
                    issues.append(f"  Paragraph {i}: unreplaced object: {run.text[:60]}")

        if issues:
            print(f"WARNING: Found {len(issues)} validation issues:")
            for issue in issues[:5]:
                print(issue)
            if len(issues) > 5:
                print(f"  ... and {len(issues) - 5} more")

    def _setup_section(self, native_section, margins: dict, footer: str, ctx: RenderContext):
        """Apply margins and footer to a native section."""
        if margins:
            native_section.left_margin = Cm(margins.get('left', 3))
            native_section.top_margin = Cm(margins.get('top', 2))
            native_section.bottom_margin = Cm(margins.get('bottom', 2.5))
            native_section.right_margin = Cm(margins.get('right', 1))

        # Unlink from previous to allow custom footer
        native_section.footer.is_linked_to_previous = False

        if footer is None or footer == "page_number":
            self._add_page_number_footer(native_section, ctx)
        elif isinstance(footer, str):
            self._add_custom_footer(native_section, footer, ctx)

    def _add_page_number_footer(self, native_section, ctx: RenderContext):
        """Add centered page number to a section's footer."""
        footer = native_section.footer
        for p in footer.paragraphs:
            p.clear()

        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

        token = self.style_sheet.get("body")
        ctx._apply_paragraph_style(p, token)
        p.alignment = 1  # CENTER (after apply_style so it sticks)
        run = p.add_run()
        ctx._apply_run_style(run, token)

        # PAGE field
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char_begin)
        run._r.append(instr)
        run._r.append(fld_char_end)

    def _add_custom_footer(self, native_section, text: str, ctx: RenderContext):
        """Add custom text to a section's footer."""
        footer = native_section.footer
        for p in footer.paragraphs:
            p.clear()

        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

        token = self.style_sheet.get("body")
        ctx._apply_paragraph_style(p, token)
        p.alignment = 1  # CENTER (after apply_style so it sticks)
        run = p.add_run(text)
        ctx._apply_run_style(run, token)

    def render_section(self, section, ctx: RenderContext):
        """Render a section — creates a new native section for VDOM sections after the first."""
        native_sections = ctx.native_doc.sections
        if self._section_counter == 0:
            native_sec = native_sections[0]
        else:
            native_sec = ctx.native_doc.add_section()

        self._setup_section(native_sec, section.margins, section.footer, ctx)
        ctx._current_section = native_sec
        self._section_counter += 1

        for child in section.children:
            child.render(self, ctx)

    def render_heading(self, heading, ctx: RenderContext):
        """Render a heading element."""
        # Get style token for this heading level
        role = f"heading_{heading.level}"
        token = self.style_sheet.get(role)

        # Apply any local overrides
        if heading.style_overrides:
            token = token.merge(**heading.style_overrides)

        # Add paragraph
        para = ctx.add_paragraph(token)

        # Set Word's built-in heading style so TOC can index it
        style_name = f'Heading {heading.level}'
        try:
            para.style = ctx._doc.styles[style_name]
        except KeyError:
            pass

        # Re-apply direct formatting (overrides the style's defaults)
        ctx._apply_paragraph_style(para, token)
        ctx.add_run(para, heading.text, token)

    def render_paragraph(self, paragraph, ctx: RenderContext):
        """Render a paragraph with inline content."""
        # Determine effective style
        if paragraph.style_role:
            token = self.style_sheet.get(paragraph.style_role)
        else:
            token = self.style_sheet.get("body")
        if paragraph.style_overrides:
            token = token.merge(**paragraph.style_overrides)

        para = ctx.add_paragraph(token)

        for item in paragraph.content:
            if isinstance(item, str):
                ctx.add_run(para, item, token)
            elif hasattr(item, 'text'):  # Span
                span_token = token.merge(**item.style_overrides)
                ctx.add_run(para, item.text, span_token)

    def render_table(self, table, ctx: RenderContext):
        """Render a table with headers and rows."""
        # Calculate column count
        col_count = 0
        if table.headers:
            col_count = len(table.headers)
        elif table.rows:
            col_count = max(len(row) for row in table.rows)

        if col_count == 0:
            return

        # Count rows
        row_count = len(table.rows)
        if table.headers:
            row_count += 1

        # Create native table
        native_table = ctx.native_doc.add_table(rows=row_count, cols=col_count)
        native_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set borders
        tbl = native_table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Remove any existing style
        tblStyle = tblPr.find(qn('w:tblStyle'))
        if tblStyle is not None:
            tblPr.remove(tblStyle)

        # Set borders
        existing_borders = tblPr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tblPr.remove(existing_borders)

        borders = OxmlElement('w:tblBorders')
        if table.visible_borders:
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                borders.append(border)
        else:
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                borders.append(border)
        tblPr.append(borders)

        # Set table width to 100%
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:type'), 'pct')
        tblW.set(qn('w:w'), '5000')

        # Fill headers
        if table.headers:
            header_token = self.style_sheet.get("table_header")
            for j, header_text in enumerate(table.headers):
                cell = native_table.cell(0, j)
                self._fill_cell(cell, header_text, header_token)

        # Fill data rows
        data_token = self.style_sheet.get("table_cell")
        for i, row_data in enumerate(table.rows):
            row_idx = i + 1 if table.headers else i
            for j, cell_data in enumerate(row_data):
                if j < col_count:
                    cell = native_table.cell(row_idx, j)
                    if isinstance(cell_data, Cell):
                        # Handle Cell object
                        token = data_token
                        if cell_data.bold:
                            token = token.merge(bold=True)
                        if cell_data.alignment:
                            token = token.merge(alignment=cell_data.alignment)
                        # Clear default paragraph and reuse it
                        for p in cell.paragraphs:
                            p.clear()
                        para = cell.paragraphs[0]
                        for elem in cell_data.content:
                            if hasattr(elem, 'content'):  # Paragraph
                                ctx._apply_paragraph_style(para, token)
                                for item in elem.content:
                                    if isinstance(item, str):
                                        ctx.add_run(para, item, token)
                                    elif hasattr(item, 'text'):  # Span
                                        span_token = token.merge(**item.style_overrides)
                                        ctx.add_run(para, item.text, span_token)
                            else:
                                self._fill_cell(cell, str(elem), token)
                    elif hasattr(cell_data, 'text'):  # Span object
                        # Handle Span object directly
                        span_token = data_token.merge(**cell_data.style_overrides)
                        cell.text = ""
                        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                        ctx._apply_paragraph_style(para, span_token)
                        ctx.add_run(para, cell_data.text, span_token)
                    else:
                        self._fill_cell(cell, str(cell_data), data_token)

    def _fill_cell(self, cell, text: str, token: StyleToken):
        """Fill a table cell with text and formatting."""
        # Clear existing content
        for p in cell.paragraphs:
            p.clear()

        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        ctx = RenderContext.__new__(RenderContext)
        ctx._doc = None
        ctx._style_sheet = self.style_sheet
        ctx._apply_paragraph_style(para, token)
        run = para.add_run(text)
        ctx._apply_run_style(run, token)

    def _embed_image(self, path: str, width_cm: float, caption: str, ctx: RenderContext):
        """Core image embedding logic — shared by Image and PlantUml elements."""
        import os
        if not os.path.exists(path):
            token = self.style_sheet.get("body")
            para = ctx.add_paragraph(token)
            ctx.add_run(para, f"[Image not found: {path}]", token)
            return

        width = Cm(width_cm) if width_cm else Cm(15)
        para = ctx.add_paragraph(self.style_sheet.get("body"))
        run = para.add_run()
        run.add_picture(path, width=width)

        if caption:
            caption_token = self.style_sheet.get("caption")
            cap_para = ctx.add_paragraph(caption_token)
            ctx.add_run(cap_para, caption, caption_token)

    def render_image(self, image, ctx: RenderContext):
        """Render an Image element."""
        self._embed_image(image.path, image.width_cm, image.caption, ctx)

    def render_plantuml(self, plantuml, ctx: RenderContext):
        """Render a PlantUML diagram — compile .puml to PNG, then embed."""
        png_path = self._compile_plantuml(plantuml.filename)
        self._embed_image(png_path, plantuml.width_cm, plantuml.caption, ctx)

    def _compile_plantuml(self, puml_path: str) -> str:
        """Compile a .puml file to PNG with content-hash caching.

        Returns path to cached PNG.
        """
        import hashlib, subprocess, tempfile, shutil, os

        with open(puml_path) as f:
            source = f.read()

        digest = hashlib.md5(source.encode()).hexdigest()
        cache_dir = os.path.join(os.path.dirname(puml_path), "cache")
        os.makedirs(cache_dir, exist_ok=True)
        png_path = os.path.join(cache_dir, f"{digest}.png")

        if os.path.exists(png_path):
            return png_path

        # Write temp .puml to compile
        with tempfile.NamedTemporaryFile(suffix=".puml", delete=False, mode="w") as f:
            f.write(source)
            tmp_puml = f.name

        try:
            subprocess.run(
                ["plantuml", "-tpng", tmp_puml],
                check=True, capture_output=True, timeout=30,
            )
            generated_png = tmp_puml.replace(".puml", ".png")
            if os.path.exists(generated_png):
                shutil.move(generated_png, png_path)
            else:
                raise RuntimeError(f"PlantUML did not produce {generated_png}")
        finally:
            os.unlink(tmp_puml)

        return png_path

    def render_code_block(self, code_block, ctx: RenderContext):
        """Render a code block."""
        token = self.style_sheet.get("code")
        for line in code_block.code.split('\n'):
            para = ctx.add_paragraph(token)
            ctx.add_run(para, line, token)

    def render_page_break(self, ctx: RenderContext):
        """Render a page break."""
        para = ctx.native_doc.add_paragraph()
        run = para.add_run()
        from docx.enum.text import WD_BREAK
        run.add_break(WD_BREAK.PAGE)

    def render_empty_line(self, empty_line, ctx: RenderContext):
        """Render empty lines."""
        if empty_line.style_role:
            token = self.style_sheet.get(empty_line.style_role)
        else:
            token = self.style_sheet.get("body")
        for _ in range(empty_line.count):
            para = ctx.add_paragraph(token)
            ctx.add_run(para, "", token)

    def render_field(self, field, ctx: RenderContext):
        """Render a dynamic field (PAGE, TOC, etc.)."""
        # Auto-update fields on document open (Word/LibreOffice)
        settings_el = ctx._doc.settings.element
        if settings_el.find(qn('w:updateFields')) is None:
            update = OxmlElement('w:updateFields')
            update.set(qn('w:val'), 'true')
            settings_el.append(update)

        token = self.style_sheet.get("body")
        para = ctx.add_paragraph(token)

        def make_run(text=None):
            run = para.add_run(text)
            ctx._apply_run_style(run, token)
            return run

        r1 = make_run()
        r1._r.append(OxmlElement('w:fldChar'))
        r1._r[-1].set(qn('w:fldCharType'), 'begin')

        r2 = make_run()
        r2._r.append(OxmlElement('w:instrText'))
        r2._r[-1].set(qn('xml:space'), 'preserve')
        r2._r[-1].text = f' {field.field_code} '

        r3 = make_run()
        r3._r.append(OxmlElement('w:fldChar'))
        r3._r[-1].set(qn('w:fldCharType'), 'separate')

        r4 = make_run('')

        r5 = make_run()
        r5._r.append(OxmlElement('w:fldChar'))
        r5._r[-1].set(qn('w:fldCharType'), 'end')

